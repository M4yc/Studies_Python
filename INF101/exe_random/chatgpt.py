from docx import Document

estoque = { 'Tomate' : [1000, 2.30],
            'Cenoura' : [500, 1.50],
            'Batata': [1500, 1.30],
            'Couve' : [300, 1.00],
            'Feijao' : [500, 6.00]}

vendas = [('Tomate', 5), ('Batata', 100), ('Cenoura', 40), ('Couve', 20), ('Feijao', 80)]

for item, (qnt,preco) in estoque.items():
    for venda_item, venda_qnt in vendas:
        if item == venda_item:
            estoque[item][0] -= venda_qnt

# Cria um documento Word
doc = Document()

# Adiciona um título
doc.add_heading('Estoque do Supermercado', 0)

# Adiciona uma tabela com cabeçalho
table = doc.add_table(rows=1, cols=3)
table.autofit = True
table.style = 'Table Grid'

table.rows[0].cells[0].text = 'Produto'
table.rows[0].cells[1].text = 'Quantidade'
table.rows[0].cells[2].text = 'Preço'

# Adiciona dados à tabela
for item, (qnt, preco) in estoque.items():
    row = table.add_row().cells
    row[0].text = item
    row[1].text = str(qnt)
    row[2].text = f'R$ {preco:.2f}'

# Salva o documento como um arquivo Word
doc.save('estoque.docx')

print("Dados do estoque exportados com sucesso para 'estoque.docx'.")
